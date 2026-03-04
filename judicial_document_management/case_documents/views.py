# case_documents/views.py
from rest_framework import viewsets, status, mixins
from rest_framework.decorators import action
from rest_framework.response import Response
from django.shortcuts import get_object_or_404
from django.contrib.contenttypes.models import ContentType
from .models import CaseDocument, DocumentTemplate
from .serializers import (
    CaseDocumentListSerializer, 
    CaseDocumentDetailSerializer,
    DocumentTemplateSerializer
)
import os
import tempfile
from django.http import FileResponse, HttpResponse
from django.conf import settings
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import status
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import mammoth
# from weasyprint import HTML
import html2text


class DocumentTemplateViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра шаблонов документов.
    Только чтение, так как управление шаблонами происходит через админку.
    """
    queryset = DocumentTemplate.objects.filter(is_active=True)
    serializer_class = DocumentTemplateSerializer
    filterset_fields = ['case_category']


class DocumentViewSet(viewsets.ModelViewSet):
    """
    ViewSet для прямого доступа к документам (не через конкретное дело).
    Используется редко, но может пригодиться для административных целей.
    """
    queryset = CaseDocument.objects.all()
    
    def get_serializer_class(self):
        if self.action == 'list':
            return CaseDocumentListSerializer
        return CaseDocumentDetailSerializer
    
    @action(detail=True, methods=['post'])
    def sign(self, request, pk=None):
        """Подписание документа"""
        document = self.get_object()
        if document.status == 'signed':
            return Response(
                {'detail': 'Документ уже подписан.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        success, message = document.sign(request.user)
        if success:
            serializer = self.get_serializer(document)
            return Response(serializer.data)
        else:
            return Response(
                {'detail': message},
                status=status.HTTP_400_BAD_REQUEST
            )

    @action(detail=False, methods=['post'], url_path='convert-word')
    def convert_word(self, request):
        """
        Конвертация Word файла в HTML для редактирования
        """
        if 'file' not in request.FILES:
            return Response(
                {'error': 'Файл не предоставлен'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        word_file = request.FILES['file']
        
        # Проверка расширения файла
        if not word_file.name.lower().endswith(('.doc', '.docx')):
            return Response(
                {'error': 'Поддерживаются только файлы .doc и .docx'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Проверка размера файла (например, не больше 10MB)
        if word_file.size > 10 * 1024 * 1024:  # 10MB
            return Response(
                {'error': 'Файл слишком большой. Максимальный размер 10MB'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Создаем временный файл
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            for chunk in word_file.chunks():
                tmp_file.write(chunk)
            tmp_path = tmp_file.name
        
        try:
            # Конвертируем Word в HTML с помощью mammoth
            with open(tmp_path, 'rb') as docx_file:
                # Настраиваем конвертацию для лучшего сохранения форматирования
                style_map = """
                p[style-name='Heading 1'] => h1:fresh
                p[style-name='Heading 2'] => h2:fresh
                p[style-name='Heading 3'] => h3:fresh
                p[style-name='Title'] => h1.title:fresh
                p[style-name='Subtitle'] => h2.subtitle:fresh
                p[style-name='Normal'] => p:fresh
                r[style-name='Strong'] => strong
                r[style-name='Emphasis'] => em
                """
                
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=style_map
                )
                html_content = result.value
                
                # Пытаемся извлечь заголовок из документа
                try:
                    from docx import Document
                    doc = Document(tmp_path)
                    title = ""
                    # Ищем первый абзац со стилем Заголовок 1 или просто первый абзац
                    for paragraph in doc.paragraphs:
                        if paragraph.style and 'Heading 1' in paragraph.style.name:
                            title = paragraph.text
                            break
                    if not title and doc.paragraphs:
                        # Берем первый непустой абзац как заголовок
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                title = paragraph.text
                                break
                except Exception as e:
                    print(f"Error extracting title: {e}")
                    title = ""
                
                return Response({
                    'content': html_content,
                    'title': title,
                    'messages': result.messages
                })
                
        except Exception as e:
            return Response(
                {'error': f'Ошибка конвертации файла: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        finally:
            # Удаляем временный файл
            try:
                os.unlink(tmp_path)
            except:
                pass

    @action(detail=True, methods=['post'], url_path='export-word')
    def export_to_word(self, request, pk=None):
        """
        Экспорт документа в Word формат
        """
        document = self.get_object()
        
        # Получаем данные из запроса или используем документ
        title = request.data.get('title', document.title)
        content = request.data.get('content', document.content)
        
        # Создаем новый Word документ
        doc = Document()
        
        # Устанавливаем стили
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        
        # Добавляем заголовок
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Конвертируем HTML в текст с сохранением форматирования
        # Используем html2text для базовой конвертации
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_emphasis = False
        h.body_width = 0  # Отключаем перенос строк
        
        text_content = h.handle(content)
        
        # Разбиваем на параграфы и добавляем в документ
        for paragraph in text_content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Сохраняем во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        # Отправляем файл пользователю
        response = FileResponse(
            open(tmp_path, 'rb'),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{title}.docx"'
        
        # Удаляем временный файл после отправки
        import atexit
        atexit.register(lambda: os.unlink(tmp_path))
        
        return response

    @action(detail=True, methods=['post'], url_path='print')
    def print_document(self, request, pk=None):
        """
        Создает версию документа для печати
        """
        document = self.get_object()
        
        # Создаем HTML для печати
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{document.title}</title>
            <style>
                body {{
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    margin: 0;
                    padding: 0;
                }}
                @media print {{
                    body {{
                        margin: 2cm;
                    }}
                    .no-print {{
                        display: none;
                    }}
                }}
                .document-header {{
                    text-align: center;
                    margin-bottom: 30px;
                }}
                .document-content {{
                    text-align: justify;
                }}
                .signature {{
                    margin-top: 50px;
                    text-align: right;
                }}
            </style>
        </head>
        <body>
            <div class="document-header">
                <h1>{document.title}</h1>
            </div>
            <div class="document-content">
                {document.content}
            </div>
            {document.signature_text and f'<div class="signature">{document.signature_text}</div>' or ''}
        </body>
        </html>
        """
        
        return Response({'html': html_content})


# Базовый класс для вложенных ViewSet'ов в производствах
class BaseCaseDocumentViewSet(mixins.CreateModelMixin,
                              mixins.RetrieveModelMixin,
                              mixins.UpdateModelMixin,
                              mixins.DestroyModelMixin,
                              mixins.ListModelMixin,
                              viewsets.GenericViewSet):
    """
    Базовый ViewSet для работы с документами внутри конкретного дела.
    Должен быть унаследован и настроен в каждом приложении производства.
    """
    serializer_class = CaseDocumentDetailSerializer
    
    def get_queryset(self):
        """Должен быть переопределен в наследниках"""
        raise NotImplementedError
    
    def get_serializer_class(self):
        if self.action == 'list':
            return CaseDocumentListSerializer
        return super().get_serializer_class()
    
    def get_serializer_context(self):
        context = super().get_serializer_context()
        # Добавляем объект дела в контекст
        context['case'] = self.get_case_object()
        return context
    
    def get_case_object(self):
        """Должен быть переопределен в наследниках"""
        raise NotImplementedError
    
    def perform_create(self, serializer):
        case = self.get_case_object()
        content_type = ContentType.objects.get_for_model(case)
        serializer.save(
            content_type=content_type,
            object_id=case.id
        )
    
    @action(detail=True, methods=['post'])
    def sign(self, request, pk=None):
        """Подписание документа"""
        document = self.get_object()
        if document.status == 'signed':
            return Response(
                {'detail': 'Документ уже подписан.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        success, message = document.sign(request.user)
        if success:
            serializer = self.get_serializer(document)
            return Response(serializer.data)
        else:
            return Response(
                {'detail': message},
                status=status.HTTP_400_BAD_REQUEST
            )


def upload_image(request):
    """
    Обработчик для загрузки изображений.
    Временная заглушка, чтобы сервер запустился.
    """
    if request.method == 'POST':
        # Здесь будет логика обработки загруженного изображения
        # Пока просто показываем сообщение об успехе
        messages.success(request, 'Изображение успешно загружено (заглушка).')
        return redirect('some-view-name')  # Перенаправьте на нужную страницу
    else:
        # Если GET-запрос, показываем форму загрузки
        return render(request, 'case_documents/upload_image.html')  # Убедитесь, что такой шаблон существует